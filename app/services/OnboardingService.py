import os
import uuid
import datetime
from docx import Document
from typing import Dict, Any, List
from config.supabase import supabase_client
from app.tools.file_handler import FileHandler
from app.tools.validators.ktp_validator import KtpValidator
from app.tools.extractors.cv_contact_extractor import CvContactExtractor
import logging

logger = logging.getLogger(__name__)

class OnboardingService:
    def __init__(self, llm):
        self.llm = llm
        self.ktp_validator = KtpValidator(llm)
        self.cv_extractor = CvContactExtractor(llm)

    async def analyze(
        self, 
        user_id: str, 
        job_position_id: int, 
        join_date: str,
        files: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Extracts data from KTP and CV and prepares data for review.
        Does NOT generate the document.
        """
        try:
            # 1. Fetch Job Details
            job_details = {"title": "", "department": ""}
            try:
                job_res = supabase_client.table("job_positions").select("title, department").eq("id", job_position_id).single().execute()
                if job_res.data:
                    job_details = job_res.data
            except Exception as e:
                logger.error(f"Error fetching job details: {str(e)}")

            # 2. Extract Data from KTP
            ktp_data = {}
            if "ktp" in files:
                try:
                    ktp_data = await self.ktp_validator(files["ktp"]["bytes"], files["ktp"]["mime_type"])
                except Exception as e:
                    logger.error(f"KTP Extraction failed: {str(e)}")

            # 3. Extract Data from CV
            cv_data = {}
            if "cv" in files:
                try:
                    cv_data = await self.cv_extractor(files["cv"]["bytes"], files["cv"]["mime_type"])
                except Exception as e:
                    logger.error(f"CV Extraction failed: {str(e)}")

            # 4. Prepare Final Data
            email_phone = f"{cv_data.get('email', '')} / {cv_data.get('phone', '')}"
            final_data = {
                "candidate_name": ktp_data.get("full_name", ""),
                "candidate_nik": ktp_data.get("nik", ""),
                "candidate_address": ktp_data.get("address", ""),
                "job_department": job_details.get("department", ""),
                "job_position": job_details.get("title", ""),
                "join_date": join_date,
                "email_and_phone_number_candidate": email_phone
            }

            # 6. Log
            input_paths = [f["path"] for f in files.values()]
            try:
                log_data = {
                    "user_id": user_id,
                    "tool_type": "onboarding_analyze",
                    "input_files": input_paths,
                    "result_json": final_data,
                    "cost_usd": 0.005,
                    "token_usage": {"prompt": 200, "completion": 100}
                }
                supabase_client.table("activity_logs").insert(log_data).execute()
            except Exception as e:
                logger.error(f"Logging failed: {str(e)}")

            return final_data

        except Exception as e:
            logger.error(f"OnboardingService Analysis Fatal Error: {str(e)}", exc_info=True)
            raise e

    async def generate_document(self, user_id: str, data: Dict[str, Any]) -> str:
        """
        Generates the DOCX file from the reviewed data.
        """
        try:
            # Fetch HR Name from Profiles
            hr_name = "Unknown HR"
            try:
                profile_res = supabase_client.table("profiles").select("full_name").eq("id", user_id).single().execute()
                if profile_res.data:
                    hr_name = profile_res.data.get("full_name", hr_name)
            except Exception as e:
                logger.error(f"Error fetching profile: {str(e)}")
            
            # Generate Date
            now = datetime.datetime.now()
            date_now_1 = now.strftime("%d %b %Y")
            
            # Fix Email/Phone Construction
            email = data.get("email", "")
            phone = data.get("phone", "")
            if email or phone:
                data["email_and_phone_number_candidate"] = f"{email} / {phone}"
            
            # Add to data for template
            data["hr_name"] = hr_name
            data["date_now_1"] = date_now_1

             # 5. Generate Document
            docx_path = self._fill_template(data)
            
            # Upload
            with open(docx_path, "rb") as f:
                docx_bytes = f.read()
            storage_path = f"generated/onboarding/{uuid.uuid4()}.docx"
            await FileHandler.upload_file(docx_bytes, storage_path, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            os.remove(docx_path)
            
            # 6. Log Generation
            try:
                log_data = {
                    "user_id": user_id,
                    "tool_type": "onboarding_generate",
                    "output_files": [storage_path],
                    "result_json": {"document_generated": True},
                    "cost_usd": 0,
                    "token_usage": {}
                }
                supabase_client.table("activity_logs").insert(log_data).execute()
            except Exception as e:
                logger.error(f"Logging failed: {str(e)}")

            return FileHandler.get_public_url(storage_path)

        except Exception as e:
            logger.error(f"OnboardingService Doc Gen Fatal Error: {str(e)}", exc_info=True)
            raise e

    def _fill_template(self, data: Dict[str, Any]) -> str:
        template_path = "assets/template/Template_Onboarding_Offboarding Form.docx"
        if not os.path.exists(template_path):
             doc = Document()
             doc.add_paragraph("Template not found.")
        else:
             doc = Document(template_path)
        
        replacements = {
            "{candidate_name}": data.get("candidate_name", ""),
            "{candidate_nik}": data.get("candidate_nik", ""),
            "{candidate_address}": data.get("candidate_address", ""),
            "{job_department}": data.get("job_department", ""),
            "{job_position}": data.get("job_position", ""),
            "{join_date}": data.get("join_date", ""),
            "{email_and_phone_number_candidate}": data.get("email_and_phone_number_candidate", ""),
            "{hr_name}": data.get("hr_name", ""),
            "{date_now_1}": data.get("date_now_1", "")
        }
        
        # Replace in Paragraphs
        for paragraph in doc.paragraphs:
            self._replace_text_in_paragraph(paragraph, replacements)
            
        # Replace in Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_text_in_paragraph(paragraph, replacements)
                        
        filename = f"/tmp/{uuid.uuid4()}.docx"
        doc.save(filename)
        return filename

    def _replace_text_in_paragraph(self, paragraph, replacements):
        if not paragraph.text:
            return
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)

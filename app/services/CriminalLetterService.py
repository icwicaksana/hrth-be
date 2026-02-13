import os
import uuid
import datetime
import tempfile
from docx import Document
from typing import Dict, Any, List
from config.supabase import supabase_client
from app.tools.file_handler import FileHandler
from app.tools.validators.ktp_validator import KtpValidator
import logging

logger = logging.getLogger(__name__)

class CriminalLetterService:
    def __init__(self, llm):
        self.llm = llm
        self.ktp_validator = KtpValidator(llm)

    async def analyze(
        self, 
        user_id: str, 
        job_position_id: int, 
        files: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Analyzes KTP to extract candidate data for the Criminal Letter.
        """
        try:
            # 1. Fetch Job Position Title
            job_position_title = "Unknown"
            try:
                job_res = supabase_client.table("job_positions").select("title").eq("id", job_position_id).single().execute()
                if job_res.data:
                    job_position_title = job_res.data["title"]
            except Exception as e:
                logger.error(f"Error fetching job position: {str(e)}")

            # 2. Extract Data from KTP
            ktp_data = {}
            if "ktp" in files:
                try:
                    ktp_data = await self.ktp_validator(files["ktp"]["bytes"], files["ktp"]["mime_type"])
                except Exception as e:
                    logger.error(f"KTP Extraction failed: {str(e)}")
            
            # 3. Calculate Age (Simple approximation)
            age = ""
            if "birth_date" in ktp_data:
                try:
                    # Expected format DD-MM-YYYY
                    birth_date_str = ktp_data["birth_date"]
                    birth_date = datetime.datetime.strptime(birth_date_str, "%d-%m-%Y")
                    today = datetime.datetime.now()
                    age = today.year - birth_date.year - ((today.month, today.day) < (birth_date.month, birth_date.day))
                    age = f"{age} Tahun"
                except Exception as e:
                    logger.warning(f"Could not calculate age: {str(e)}")

            # 4. Prepare Final Data
            now = datetime.datetime.now()
            final_data = {
                "candidate_name": ktp_data.get("full_name", ""),
                "candidate_gender": ktp_data.get("gender", ""),
                "birth_location": ktp_data.get("birth_place", ""),
                "birth_date": ktp_data.get("birth_date", ""),
                "candidate_age": str(age),
                "job_position": job_position_title,
                "candidate_address": ktp_data.get("address", ""),
                "date_now_2": now.strftime("%d %B %Y")
            }

            # 5. Log Analysis
            input_paths = [f["path"] for f in files.values()]
            try:
                log_data = {
                    "user_id": user_id,
                    "tool_type": "criminal_letter_analyze",
                    "input_files": input_paths,
                    "result_json": final_data,
                    "cost_usd": 0.002, 
                    "token_usage": {"prompt": 150, "completion": 80} 
                }
                supabase_client.table("activity_logs").insert(log_data).execute()
            except Exception as e:
                logger.error(f"Logging failed: {str(e)}")

            return final_data

        except Exception as e:
            logger.error(f"CriminalLetterService Analysis Fatal Error: {str(e)}", exc_info=True)
            raise e

    async def generate_document(self, user_id: str, data: Dict[str, Any]) -> str:
        """
        Generates the 'Surat Keterangan Bebas Tindakan Kriminal' DOCX.
        """
        try:
            # Generate Document
            docx_path = self._fill_template(data)
            
            # Upload
            with open(docx_path, "rb") as f:
                docx_bytes = f.read()
            storage_path = f"generated/criminal_letter/{uuid.uuid4()}.docx"
            await FileHandler.upload_file(docx_bytes, storage_path, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            os.remove(docx_path)
            
            # Log Generation
            try:
                log_data = {
                    "user_id": user_id,
                    "tool_type": "criminal_letter_generate",
                    "output_files": [storage_path],
                    "result_json": {"document_generated": True},
                    "cost_usd": 0,
                    "token_usage": {}
                }
                supabase_client.table("activity_logs").insert(log_data).execute()
            except Exception as e:
                logger.error(f"Logging failed: {str(e)}")

            return FileHandler.get_public_url(storage_path)

        except Exception as e:
            logger.error(f"CriminalLetterService Doc Gen Fatal Error: {str(e)}", exc_info=True)
            raise e

    def _fill_template(self, data: Dict[str, Any]) -> str:
        template_path = "assets/template/Template_Surat Keterangan Bebas Tindakan Kriminal.docx"
        if not os.path.exists(template_path):
             doc = Document()
             doc.add_paragraph("Template not found.")
        else:
             doc = Document(template_path)
        
        replacements = {
            "{candidate_name}": data["candidate_name"],
            "{candidate_gender}": data["candidate_gender"],
            "{birth_location}": data["birth_location"],
            "{birth_date}": data["birth_date"],
            "{candidate_age}": data["candidate_age"],
            "{job_position}": data["job_position"],
            "{candidate_address}": data["candidate_address"],
            "{date_now_2}": data["date_now_2"]
        }
        
        # Replace in Paragraphs
        for paragraph in doc.paragraphs:
            self._replace_text_in_paragraph(paragraph, replacements)
            
        # Replace in Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_text_in_paragraph(paragraph, replacements)
                        
        # Save - using cross-platform temp directory
        temp_dir = tempfile.gettempdir()
        filename = os.path.join(temp_dir, f"{uuid.uuid4()}.docx")
        doc.save(filename)
        return filename

    def _replace_text_in_paragraph(self, paragraph, replacements):
        if not paragraph.text:
            return
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)


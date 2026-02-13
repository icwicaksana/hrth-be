import os
import uuid
import datetime
import tempfile
from docx import Document
from typing import Dict, Any, List
from config.supabase import supabase_client
from app.tools.file_handler import FileHandler
from app.tools.validators.ktp_validator import KtpValidator
from app.tools.validators.academic_validator import AcademicValidator
from app.tools.validators.criminal_validator import CriminalValidator
import logging

logger = logging.getLogger(__name__)

class BgCheckService:
    def __init__(self, llm):
        self.llm = llm
        self.ktp_validator = KtpValidator(llm)
        self.academic_validator = AcademicValidator(llm)
        self.criminal_validator = CriminalValidator(llm)

    async def analyze(self, user_id: str, manual_data: Dict[str, Any], files: Dict[str, Any]) -> Dict[str, Any]:
        """
        Analyzes uploaded files using LLM validators and prepares data for review.
        Does NOT generate the document.
        """
        try:
            # 0. Fetch Job Position Title
            job_position_title = "Unknown"
            if "job_position_id" in manual_data:
                try:
                    job_res = supabase_client.table("job_positions").select("title").eq("id", manual_data["job_position_id"]).single().execute()
                    if job_res.data:
                        job_position_title = job_res.data["title"]
                except Exception as e:
                    logger.error(f"Error fetching job position: {str(e)}")

            # Fetch HR Name from Profiles
            hr_name = "Unknown HR"
            try:
                profile_res = supabase_client.table("profiles").select("full_name").eq("id", user_id).single().execute()
                if profile_res.data:
                    hr_name = profile_res.data.get("full_name", hr_name)
            except Exception as e:
                logger.error(f"Error fetching profile: {str(e)}")
            
            # 1. KTP Validation
            ktp_result = {"is_valid": False, "reasoning": "No KTP file"}
            if "ktp" in files:
                try:
                    ktp_result = await self.ktp_validator(files["ktp"]["bytes"], files["ktp"]["mime_type"])
                except Exception as e:
                    logger.error(f"KTP Validator failed: {str(e)}")
                    ktp_result = {"is_valid": False, "reasoning": f"Validator Error: {str(e)}"}
            
            # 2. Academic Validation
            academic_result = {"is_valid": False, "reasoning": "No Academic file"}
            if "academic" in files:
                try:
                    academic_result = await self.academic_validator(files["academic"]["bytes"], files["academic"]["mime_type"])
                except Exception as e:
                    logger.error(f"Academic Validator failed: {str(e)}")
                    academic_result = {"is_valid": False, "reasoning": f"Validator Error: {str(e)}"}
                
            # 3. Criminal Validation
            criminal_result = {"is_valid": False, "reasoning": "No Criminal file"}
            if "criminal" in files:
                try:
                    # We need name from KTP or manual input to cross-reference
                    candidate_name = ktp_result.get("full_name") or manual_data.get("full_name") or "Unknown"
                    criminal_result = await self.criminal_validator(files["criminal"]["bytes"], candidate_name, files["criminal"]["mime_type"])
                except Exception as e:
                    logger.error(f"Criminal Validator failed: {str(e)}")
                    criminal_result = {"is_valid": False, "reasoning": f"Validator Error: {str(e)}"}
            
            # 4. Aggregate Results
            now = datetime.datetime.now()
            final_data = {
                "manual": manual_data,
                "job_position": job_position_title,
                "hr_name": hr_name,
                "ktp": ktp_result,
                "academic": academic_result,
                "criminal": criminal_result,
                "date_now_1": now.strftime("%d %b %Y"), # e.g. 13 Dec 2025
                "date_now_2": now.strftime("%d %B %Y")  # e.g. 13 December 2025
            }

            # 6. Log Analysis Activity
            input_paths = [f["path"] for f in files.values()]
            try:
                log_data = {
                    "user_id": user_id,
                    "tool_type": "bg_check_analyze",
                    "input_files": input_paths,
                    "result_json": final_data,
                    "cost_usd": 0.005, # Estimate
                    "token_usage": {"prompt": 300, "completion": 150} # Estimate
                }
                supabase_client.table("activity_logs").insert(log_data).execute()
            except Exception as e:
                logger.error(f"Logging failed: {str(e)}")

            return final_data

        except Exception as e:
            logger.error(f"BgCheckService Analysis Fatal Error: {str(e)}", exc_info=True)
            raise e

    async def generate_document(self, user_id: str, data: Dict[str, Any]) -> str:
        """
        Generates the DOCX file from the reviewed data.
        """
        try:
            # Generate Document (Filling Template)
            docx_path = self._fill_template(data)
            
            # Upload Docx
            with open(docx_path, "rb") as f:
                docx_bytes = f.read()
            storage_path = f"generated/bg_check/{uuid.uuid4()}.docx"
            await FileHandler.upload_file(docx_bytes, storage_path, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            os.remove(docx_path)

            # Log Generation Activity
            try:
                log_data = {
                    "user_id": user_id,
                    "tool_type": "bg_check_generate",
                    "output_files": [storage_path],
                    "result_json": {"document_generated": True},
                    "cost_usd": 0, 
                    "token_usage": {} 
                }
                supabase_client.table("activity_logs").insert(log_data).execute()
            except Exception as e:
                logger.error(f"Logging failed: {str(e)}")
            
            return FileHandler.get_public_url(storage_path)

        except Exception as e:
            logger.error(f"BgCheckService Doc Gen Fatal Error: {str(e)}", exc_info=True)
            raise e

    def _fill_template(self, data: Dict[str, Any]) -> str:
        # Load Template
        template_path = "assets/template/Template_Formulir Pelaksanaan Background Screening.docx"
        if not os.path.exists(template_path):
             doc = Document()
             doc.add_paragraph("Template not found, generated blank.")
        else:
             doc = Document(template_path)
        
        manual = data["manual"]
        ktp = data["ktp"]
        
        # Prepare Values
        replacements = {
            "{candidate_name}": ktp.get("full_name", "") or "",
            "{candidate_nik}": ktp.get("nik", "") or "",
            "{job_position}": data["job_position"] or "",
            "{date_now_1}": data["date_now_1"] or "",
            "{date_now_2}": data["date_now_2"] or "",
            "{hr_name}": data["hr_name"] or "Unknown HR",
            "{ktp_check}": "v" if ktp.get("is_valid") else "",
            "{interview_check}": "v" if manual["interview_passed"] else "",
            "{academic_check}": "v" if data["academic"].get("is_valid") else "",
            "{reference_check}": "v" if manual["reference_checked"] else "",
            "{criminal_check}": "v" if data["criminal"].get("is_valid") else "",
        }
        
        # Replace in Paragraphs
        for paragraph in doc.paragraphs:
            self._replace_text_in_paragraph(paragraph, replacements)
            
        # Replace in Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_text_in_paragraph(paragraph, replacements)
                
        # Save - using cross-platform temp directory
        temp_dir = tempfile.gettempdir()
        filename = os.path.join(temp_dir, f"{uuid.uuid4()}.docx")
        doc.save(filename)
        return filename

    def _replace_text_in_paragraph(self, paragraph, replacements):
        if not paragraph.text:
            return
        
        # Iterate over all replacements
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)
